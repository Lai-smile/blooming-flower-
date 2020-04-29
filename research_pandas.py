import pandas as pd
import numpy as np

# data_df = pd.DataFrame(
#         np.arange(6315).reshape(1263, 5),
#         columns=[
#             '客户代码',
#             '原始订单名',
#             '原始文件名',
#             '要提取的文件名',
#             '提取项',
#         ]
#     )
# print(data_df.describe())
# print(data_df.count())
# writer = pd.ExcelWriter('play_data.xlsx')
# data_df.to_excel(writer, float_format='%.2f')
# writer.save()
import os
# str1='(0.09平米) 12-31 9 8B07405BA0.ZIP'

# print(str1.split('.')[0].split(' ')[-1])
# print(str1.split(' '))
# print(str1.split('.')[1])
# print(str1.split('.')[1].split(' ')[-1])
# print(str1.replace( 'ZIP', 'rar'))

# basic_url=r'C:\Users\Administrator\Desktop\bo74-PART'
# v='(0.002平米) 01-14 2 2B0741Q2A0/'
# print(os.listdir(os.path.join(basic_url,v))[-1])
# print(os.path.join(basic_url,v).split('/')[0])

# w = wc.Dispatch('Word.Application')
# document = w.Documents.Open(os.path.join(full_path, original_file))  # 读入文件
# tables = document.tables  # 获取文件中的表格集
# table = tables[0]  # 获取文件中的第一个表格
# if table.Rows[5].Cells[0].Range.text==True:
#     WorksRequirements_0 =table.Rows[5].Cells[0].Range.text.strip()  # 工程要求或技术要求内容
#     WorksRequirements_1 = table.Rows[5].Cells[0].Range.text.strip()  # 工程要求或技术要求内容
#     WorksRequirements_2 = table.Rows[5].Cells[1].Range.text.strip()  # 工程要求或技术要求内容
#     specialThings_0 = table.Rows[6].Cells[0].Range.text  # 注意事项
#     specialThings_1 = table.Rows[6].Cells[1].Range.text  # 注意事项内容
#
#     content = {WorksRequirements_0: WorksRequirements_2, specialThings_0: specialThings_1}
#     content_dict = {original_file: content}
#     content_list.append(content_dict)
#     content_dict = {original_file: '该文档无相应技术要求'}
#     content_list.append(content_dict)
# elif table.Rows[5].Cells[0].Range.text==True and table.Rows[6].Cells[0].Range.text==True:
#     WorksRequirements_0 =table.Rows[5].Cells[0].Range.text.strip()  # 工程要求或技术要求内容
#     WorksRequirements_1 = table.Rows[5].Cells[0].Range.text.strip()  # 工程要求或技术要求内容
#     WorksRequirements_2 = table.Rows[5].Cells[1].Range.text.strip()  # 工程要求或技术要求内容
#     specialThings_0 = table.Rows[6].Cells[0].Range.text  # 注意事项
#     specialThings_1 = table.Rows[6].Cells[1].Range.text  # 注意事项内容
#
#     content = {WorksRequirements_0: WorksRequirements_2, specialThings_0: specialThings_1}
#     content_dict = {original_file: content}
#     content_list.append(content_dict)
#     content_dict = {original_file: '该文档无相应技术要求'}
#     content_list.append(content_dict)
# else:
#     WorksRequirements_0 = '不是要提取的内容'
#     WorksRequirements_1 = '不是要提取的内容'
#     WorksRequirements_2 = '不是要提取的内容'
#     specialThings_0 = table.Rows[6].Cells[0].Range.text  # 注意事项
#     specialThings_1 = table.Rows[6].Cells[1].Range.text  # 注意事项内容
#     content = {WorksRequirements_0: WorksRequirements_2, specialThings_0: specialThings_1}
#     content_dict = {original_file: content}
#     content_list.append(content_dict)
#     content_dict = {original_file: '该文档无相应技术要求'}
#     content_list.append(content_dict)
# w.Documents.Close()
# w.Quit()
# sdkj={}
# jf={}
# dlkl={}
# fja=[]
# fja.append(sdkj)
# fja.append(jf)
# fja.append(dlkl)
# print(fja)
import os
import patoolib

base_url = "F:\B074"
unzip_high_url = r"F:\B074_originalI_file"
#列出压缩包内的文件
# patoolib.list_archive("package.deb")
# def decompress_file(filename):
    #先判断文件类型
    # if filename.endswith('zip') or filename.endswith('ZIP') or filename.endswith('rar'):
        # 获取文件名的绝对路径
# patoolib.extract_archive(r'C:\Users\Administrator\Desktop\(0.02平米) 09-19 2 8B07404FA0.zip',unzip_high_url)

# print(bsse_url.split('.')[0].split('-')[-1])
# patoolib.extract_archive("archive.zip", outdir="/tmp")
# dirs=os.listdir(base_url)
# for dir in dirs:
#     print(dir)
# dirs = os.path.split(bsse_url)
# print(dirs)
# print(os.path.split(bsse_url))
# FS='FP7.820.12308.pdf'
# print(FS.split('.')[0].startswith('F'))
# text_num=3
# content_title_list=[]
# for num in range(text_num):
#     content_title_list.append('技术要求')
# print(content_title_list)
import docx
from docx import Document
import win32com.client
# doc2docx(bsse_url)
bssc_url = "M_FP7.820.170.docx"
bsse_url = r"F:\b074_10\2B0741GHA0\(0.002平米) 08-23 1 2B0741GHA0\B074内单模板（新单模板）.doc"
# print(os.path.abspath(bsse_url))
# print(Document(bsse_url, bssc_url))
# di=Document(bsse_url)
# tables=di.tables
# table=tables[0]
# for i in range(1,len(table.rows)):
# gc=table.cell(5,2).text
# print(gc)
w=win32com.client.constants
word = win32com.client.gencache.EnsureDispatch('kWord.Application')
doc=word.Documents.Open(bsse_url)
t=doc.Tables[0]
ws=t.Rows[5].Cell[1].Range.Text
print(ws)
doc.close()