import os
from file_utilities import unpack_zip
from tokenization.get_text_from_file import get_word_tokens
from tokenization.pdf_tokenization.get_table_from_pdf import get_pdf_table
import pandas as pd
import numpy as np
import docx
import win32com.client
from docx.document import Document

def doc2docx(doc_in):
    docx_out = doc_in[:doc_in.rfind('.')] + ".docx"
    app = win32com.client.Dispatch('Word.Application')
    app.Visible = 0
    app.DisplayAlerts = 0
    doc_in = app.Documents.Open(doc_in)
    doc_in.SaveAs(docx_out, 16)
    # 16 for docx
    doc_in.Close()
    app.Quit()

    return docx_out


def decompress():
    base_url = r"F:\b074_10"
    original_order_list = []    #原始订单名列表
    original_file_list = [] #原始文件名列表
    article_name_list = [] #文档名列表
    content_title_list = [] #提取项（工程要求或技术要求）
    content_list = []   #内容列表

    dirs = os.listdir(base_url)
    for dir in dirs:
        original_order_name = dir
        # original_order_list.append(original_order_name)
        content = []
        try:
            ewr = unpack_zip(dir, base_url)
            if not ewr:
                pass
            else:
                order_file = ewr.split('/')[0]
                if order_file:  # 解压原始订单名压缩包的文件
                    folder_list = os.listdir(order_file)
                    for file in folder_list:
                        if not file:
                            pass
                        else:
                            if file.split(' ')[-1] == dir.split('.')[0] or file.split(' ')[-1] == \
                                    dir.split('.')[0].split('-')[-1]:
                                file_path = os.path.join(order_file, file)
                                original_subfile_list=os.listdir(os.path.join(order_file, file))
                                print(original_subfile_list)
                                for inner_file in original_subfile_list:
                                    # original_file_name = file + '.zip'  # 原始订单名
                                    # original_file_list.append(original_file_name)
                                    doc_dirname_path = os.path.join(file_path, inner_file)
                                    if inner_file.endswith('doc'):
                                        # doc2docx(doc_dirname_path)
                                        doc_dirname_path_to=doc_dirname_path[:doc_dirname_path.rfind('.')]+'.docx'
                                        w=win32com.client.DispatchEx('Word.Application')
                                        doc_ducu=w.Documents.Open(doc_dirname_path)
                                        doc_ducu.SaveAs(doc_dirname_path_to,16)
                                        document_file=docx.Document(doc_dirname_path_to)
                                        print("段落数:" + str(len(document_file.paragraphs)))
                                        paragraphs_num=len(document_file.paragraphs)
                                        # 输出每一段的内容
                                        for para in document_file.paragraphs:
                                            print(para)

                                        # try:
                                        # print(doc_dirname_path)
                                        # large = get_word_tokens(doc_dirname_path)
                                        #
                                        # if large:
                                        #     large_list = large[0][5][1]
                                        #     for i in range(len(large_list)):
                                        #         if large_list[i][-1] == '工程要求':

                                                    # WorksRequirements_content = large_list[i + 1][-1]
                                                    # row_1=re.search(r'^([\u4e00-\u9fa5]*).*,$', WorksRequirements_content)
                                                    # content_list.append(row_1)
                                                    # original_order_list.append(original_order_name)
                                                    # original_file_name = file + '.zip'  # 原始订单名
                                                    # doc_name = inner_file
                                                    # original_file_list.append(original_file_name)
                                                    # article_name_list.append(doc_name)
                                                    # content_title_list.append('工程要求')
                                                    # row_2=re.search(r'^([\u4e00-\u9fa5]*).*([\u4e00-\u9fa5]*)$', WorksRequirements_content)
                                                    # content_list.append(row_2)
                                                    # original_order_list.append(original_order_name)
                                                    # original_file_name = file + '.zip'  # 原始订单名
                                                    # doc_name = inner_file
                                                    # original_file_list.append(original_file_name)
                                                    # article_name_list.append(doc_name)
                                                    # content_title_list.append('工程要求')
                                                    # row_3=re.search(r'^([0-9]*).*([0-9]*)$', WorksRequirements_content)
                                                    # content_list.append(row_3)
                                                    # original_order_list.append(original_order_name)
                                                    # original_file_name = file + '.zip'  # 原始订单名
                                                    # doc_name = inner_file
                                                    # original_file_list.append(original_file_name)
                                                    # article_name_list.append(doc_name)
                                                    # content_title_list.append('工程要求')
                                                    # row_4=re.search(r'^([\u4e00-\u9fa5]*).*。$', WorksRequirements_content)
                                                    # content_list.append(row_4)
                                                    # original_order_list.append(original_order_name)
                                                    # original_file_name = file + '.zip'  # 原始订单名
                                                    # doc_name = inner_file
                                                    # original_file_list.append(original_file_name)
                                                    # article_name_list.append(doc_name)
                                                    # content_title_list.append('工程要求')
                                                    # row_5=re.search(r'^([\u4e00-\u9fa5]*).*\)$', WorksRequirements_content)
                                            content_list.append(para)
                                            original_order_list.append(original_order_name)
                                            original_file_name = file + '.zip'  # 原始订单名
                                            doc_name = inner_file
                                            original_file_list.append(original_file_name)
                                            article_name_list.append(doc_name)
                                            content_title_list.append('工程要求')
                                            specialRequirements_content = para[-1][-1]
                                            content_list.append(specialRequirements_content)
                                            original_order_list.append(original_order_name)
                                            original_file_name = file + '.zip'  # 原始订单名
                                            doc_name = inner_file
                                            original_file_list.append(original_file_name)
                                            article_name_list.append(doc_name)
                                            content_title_list.append('注意事项')

                                    if os.path.isdir(doc_dirname_path)==True:   # 如果与doc文件同目录的另外一个文件是目录
                                        innerest_file_list = os.listdir(doc_dirname_path)  # 最里层文件
                                        for innerest_file in innerest_file_list:
                                            if innerest_file.endswith('pdf'):
                                                pdf_path = os.path.join(doc_dirname_path, innerest_file)  # pdf路径
                                                try:
                                                    extract_pdf = get_pdf_table(pdf_path)
                                                    if not extract_pdf:
                                                        pass
                                                    else:
                                                        pdf_name=innerest_file
                                                        for content_num in range(len(extract_pdf)):
                                                            text_start_str=extract_pdf[content_num][5][-1].split('.')[0]
                                                            if text_start_str.isdigit():
                                                                tecnologhsquirements_content=extract_pdf[content_num][5][-1]
                                                                content_list.append(tecnologhsquirements_content)
                                                                original_file_name = file + '.zip'
                                                                content_title_list.append('技术要求')
                                                                original_order_list.append(original_order_name)
                                                                article_name_list.append(pdf_name)
                                                                original_file_list.append(original_file_name)

                                                except Exception as e:
                                                    print(e)

        except Exception as e:
            print(e)

    print(len(original_order_list))
    print(len(original_file_list))
    print(len(article_name_list))
    print(len(content_title_list))
    print(len(content_list))

    print(original_order_list)
    print(original_file_list)
    print(article_name_list)
    print(content_title_list)
    print(content_list)
    bo74_df = pd.DataFrame(np.arange(len(content_list) * 4).reshape(len(content_list), 4))
    bo74_df['原始订单名'] = original_order_list
    bo74_df['原始文件名'] = original_file_list
    bo74_df['doc文件名'] = article_name_list
    bo74_df['提取项'] = content_title_list
    bo74_df['提取内容'] = content_list
    # bo74_df.insert(0, '原始订单名', list(original_order_file_dict.keys()))
    # bo74_df.insert(1, '原始文件名', list(original_file_doc_dict.keys()))
    # bo74_df.insert(2, 'doc文件名', list(content_dict.keys()))
    # bo74_df.insert(3, '提取项', content_dict.values())
    bo74_df.to_excel('datas_of_B074.xlsx')


if __name__ == "__main__":
    # run_bo74()
    decompress()
